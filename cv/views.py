import os
from django.shortcuts import render
from django.http import HttpResponse
from django.core.mail import EmailMessage
from rest_framework.permissions import IsAuthenticated
from rest_framework.generics import CreateAPIView,RetrieveUpdateDestroyAPIView,ListAPIView,RetrieveAPIView, UpdateAPIView
from rest_framework.views import APIView
from .serializers import EducationSerializer,JuniorCertTestSerializer,ExperienceSerializer,ReferenceSerializer,CvSerializer, SkillSerializer,QualitiesSerializer, LeavingCertTestSerializer, StudentSerializer, InterestSerializer, AdditionalInfoSerializer
from .models import CV,Education,JuniorCertTest,Experience,Reference,JobTitle,Qualities,Skills,LeavingCertTest, Interests, AdditionalInfo
from users.models import Student
from django.template.loader import render_to_string
from weasyprint import HTML
from rest_framework.exceptions import  ValidationError
from rest_framework import status
from rest_framework.response import Response
from reportlab.lib.pagesizes import A4
from io import BytesIO
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Frame
from reportlab.platypus import Table, TableStyle
from docx import Document
from docx.shared import Pt, RGBColor


class AdditionalInfoViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = AdditionalInfoSerializer
    queryset = AdditionalInfo.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            add_info=AdditionalInfo.objects.filter(user=student.student)
            serializer = AdditionalInfoSerializer(add_info,many=True)
            return Response(serializer.data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        try:
            add_info_serializer_obj=AdditionalInfoSerializer(instance='',data=request.data,many=True, context=request)

            if add_info_serializer_obj.is_valid(raise_exception=True):
                    add_info_serializer_obj.save()
                    return Response(add_info_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(add_info_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
    def delete(self, request, pk):
        try:
            add_info = AdditionalInfo.objects.get(pk=pk)
            add_info.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class AdditionalInfoUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = AdditionalInfoSerializer
    queryset = AdditionalInfo.objects.all()



class CvViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = CvSerializer
    queryset=CV.objects.all()
    lookup_field = 'id'

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            cv=CV.objects.filter(user=student.student)
            serializer = CvSerializer(cv, many=True)
            return Response(serializer.data)
    
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    
    def post(self, request, *args, **kwargs):
        try:
            cv_serializer_obj=CvSerializer(instance='',data=request.data,many=True, context=request)
            try:
                email = request.data[0].get('email')
                cv_serializer_obj.validate_email(email)
            except ValidationError as e:
                error_message = str(e.detail[0]) if isinstance(e.detail, list) else str(e.detail)
                return Response({"message": error_message}, status=400)

            if cv_serializer_obj.is_valid(raise_exception=True):
                    cv_serializer_obj.save()
                    student = self.request.user.student
                    if student.cv_completed is not True:
                        student.current_step = 1
                        student.save()
                    return Response(cv_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(cv_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            error_message = str(e.detail[0]) if isinstance(e.detail, list) else str(e.detail)
            return Response({"message": error_message}, status=400)

class CVUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = CvSerializer
    queryset = CV.objects.all()


class EducationViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = EducationSerializer
    queryset = Education.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=Education.objects.filter(user=student.student)
            junior=JuniorCertTest.objects.filter(user=student.student)
            leaving = LeavingCertTest.objects.filter(user=student.student)
            serializer = EducationSerializer(edu, many=True)
            serializer2 = JuniorCertTestSerializer(junior, many=True)
            serializer3 = LeavingCertTestSerializer(leaving, many=True)
            data = {
            "education_data": serializer.data,
            "junior_data": serializer2.data,
            "leaving_data": serializer3.data,
        }
            return Response(data)

        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        if len(request.data.get('junior_data')) == 0:
            JuniorCertTest.objects.filter(user=request.user.student).delete()
        if len(request.data.get('leaving_data')) == 0:
            LeavingCertTest.objects.filter(user=request.user.student).delete()
        try:
            education_serializer_obj=EducationSerializer(instance='',data=request.data.get('education_data'),many=True, context=request)
            junior_serializer_obj=JuniorCertTestSerializer(instance='',data=request.data.get('junior_data'),many=True, context=request)
            leaving_serializer_obj=LeavingCertTestSerializer(instance='',data=request.data.get('leaving_data'),many=True, context=request)

            try:
                enddate = request.data.get('education_data')[0].get('enddate')
                education_serializer_obj.validate_enddate(enddate)
            except ValidationError as e:
                error_message = str(e.detail[0]) if isinstance(e.detail, list) else str(e.detail)
                return Response({"message": error_message}, status=400)

            if education_serializer_obj.is_valid(raise_exception=True):
                if junior_serializer_obj.is_valid(raise_exception=True):
                    if leaving_serializer_obj.is_valid(raise_exception=True):
                        education_serializer_obj.save()
                        junior_serializer_obj.save()
                        leaving_serializer_obj.save()
                        data={
                            "education_data": education_serializer_obj.data,
                            "junior_data":junior_serializer_obj.data,
                            "leaving_data":leaving_serializer_obj.data
                        }
                        student = self.request.user.student
                        if student.cv_completed is not True:
                            student.current_step = 2
                            student.save()
                        return Response(data, status=status.HTTP_201_CREATED)
            else:
                return Response(education_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            error_message = str(e.detail[0]) if isinstance(e.detail, list) else str(e.detail)
            return Response({"message": error_message}, status=400)
    
    def delete(self, request, pk):
        """Delete Education"""
        try:
            education = Education.objects.get(pk=pk)
            education.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class EducationViewUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = EducationSerializer
    queryset = Education.objects.all()

class JuniorCertTestViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = JuniorCertTestSerializer
    queryset = JuniorCertTest.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=JuniorCertTest.objects.filter(user=student.student).last()
            serializer = JuniorCertTestSerializer(edu)
            return Response(serializer.data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def create(self, request, *args, **kwargs):
        try:
            many = isinstance(request.data, list)
            serializer = self.get_serializer(data=request.data, many=many)
            serializer.is_valid(raise_exception=True)
            self.perform_create(serializer)
            headers = self.get_success_headers(serializer.data)
            return Response(serializer.data, headers=headers)
        except Exception as e:
            raise e
    
    def delete(self, request, pk):
        """Delete Junior"""
        try:
            education = JuniorCertTest.objects.get(pk=pk)
            education.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)


class JuniorViewUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = JuniorCertTestSerializer
    queryset = JuniorCertTest.objects.all()


class LeavingCertTestViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = LeavingCertTestSerializer
    queryset = LeavingCertTest.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            leaving_data=LeavingCertTest.objects.filter(user=student.student)
            serializer = LeavingCertTestSerializer(leaving_data,many=True)
            return Response(serializer.data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        try:
            leaving_data_serializer_obj=LeavingCertTestSerializer(instance='',data=request.data,many=True, context=request)

            if leaving_data_serializer_obj.is_valid(raise_exception=True):
                    leaving_data_serializer_obj.save()
                    return Response(leaving_data_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(leaving_data_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
    def delete(self, request, pk):
        try:
            leaving_data = LeavingCertTest.objects.get(pk=pk)
            leaving_data.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class LeavingViewUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = LeavingCertTestSerializer
    queryset = LeavingCertTest.objects.all()

class ExperienceViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ExperienceSerializer

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=Experience.objects.filter(user=student.student)
            serializer = ExperienceSerializer(edu, many=True)
            return Response(serializer.data)
        except Exception as e:
           error_message = next(iter(e.detail.values()))[0] if isinstance(e.detail, dict) else str(e.detail[0])
           return Response({'message': str(error_message)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        try:
            experience_serializer_obj=ExperienceSerializer(instance='',data=request.data,many=True, context=request)

            if experience_serializer_obj.is_valid(raise_exception=True):
                    experience_serializer_obj.save()
                    student = self.request.user.student
                    if student.cv_completed is not True:
                        student.current_step = 3
                        student.save()
                    return Response(experience_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(experience_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           error_message = next(iter(e.detail.values()))[0] if isinstance(e.detail, dict) else str(e.detail[0])
           return Response({'message': str(error_message)}, status=status.HTTP_400_BAD_REQUEST)
    
    def delete(self, request, pk):
        """Delete Experience"""
        try:
            experience = Experience.objects.get(pk=pk)
            experience.delete()
            return Response(status=status.HTTP_204_NO_CONTENT)
        except Exception as e:
            return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class ExperienceViewUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ExperienceSerializer
    queryset = Experience.objects.all()


class ReferenceViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ReferenceSerializer
    queryset = Reference.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=Reference.objects.filter(user=student.student)
            serializer = ReferenceSerializer(edu, many=True)
            return Response(serializer.data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        try:
            refer_serializer_obj=ReferenceSerializer(instance='',data=request.data,many=True, context=request)

            if refer_serializer_obj.is_valid(raise_exception=True):
                    refer_serializer_obj.save()
                    student = self.request.user.student
                    if student.cv_completed is not True:
                        student.current_step = 6
                        student.cv_completed = True
                        student.save()
                    return Response(refer_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(refer_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class ReferenceViewUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ReferenceSerializer
    queryset = Reference.objects.all()


class SkillsViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = SkillSerializer
    queryset = Skills.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=Skills.objects.filter(user=student.student)
            quality=Qualities.objects.filter(user=student.student)
            serializer = SkillSerializer(edu,many=True)
            serializer2= QualitiesSerializer(quality, many =True)
            data={
                'skill_data':serializer.data,
                'quality_data':serializer2.data,
            }
            return Response(data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        student =self.request.user
        Skills.objects.filter(user=student.student).delete()
        Qualities.objects.filter(user=student.student).delete()
        try:
            skill_serializer_obj=SkillSerializer(instance='',data=request.data.get('skill_data'),many=True, context=request)
            quality_serializer_obj=QualitiesSerializer(instance='',data=request.data.get('quality_data'),many=True, context=request)
            print("workinngggg")
            if skill_serializer_obj.is_valid(raise_exception=True):
                print("worror")
                skill_serializer_obj.save()
                if quality_serializer_obj.is_valid(raise_exception=True):
                    print("worror222")
                    quality_serializer_obj.save()
                    data={
                        "skill_data": skill_serializer_obj.data,
                        "quality_data":quality_serializer_obj.data
                    }
                    student = self.request.user.student
                    if student.cv_completed is not True:
                        print("workroor44")
                        student.current_step = 4
                        student.save()     
                    return Response(data, status=status.HTTP_201_CREATED)
            else:
                return Response(skill_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class SkillsUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = SkillSerializer
    queryset = Skills.objects.all()


class QualityViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = QualitiesSerializer
    queryset = Qualities.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=Qualities.objects.filter(user=student.student)
            serializer = QualitiesSerializer(edu,many=True)
            return Response(serializer.data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        try:
            quality_serializer_obj=QualitiesSerializer(instance='',data=request.data,many=True, context=request)

            if quality_serializer_obj.is_valid(raise_exception=True):
                    quality_serializer_obj.save()
                    return Response(quality_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(quality_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class QualityUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = QualitiesSerializer
    queryset = Qualities.objects.all()

class InterestViewRelated(CreateAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = InterestSerializer
    queryset = Interests.objects.all()

    def get(self, request):
        """Fetch All Chocies"""
        try:
            student =self.request.user
            edu=Interests.objects.filter(user=student.student)
            serializer = InterestSerializer(edu,many=True)
            return Response(serializer.data)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    def post(self, request, *args, **kwargs):
        try:
            interest_serializer_obj=InterestSerializer(instance='',data=request.data,many=True, context=request)

            if interest_serializer_obj.is_valid(raise_exception=True):
                    interest_serializer_obj.save()
                    return Response(interest_serializer_obj.data, status=status.HTTP_201_CREATED)
            else:
                return Response(interest_serializer_obj.errors, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
           return Response({'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class InterestUpdate(RetrieveUpdateDestroyAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = InterestSerializer
    queryset = Interests.objects.all()

# def create_pdf_with_content(user_obj,cv_obj, education_obj,junioir_cert_obj, leave_cert_obj, exp_obj, skill_obj, quality_obj,refer_obj):
#     # Create a buffer to store the PDF content
#     buffer = BytesIO()
#     # Create a SimpleDocTemplate with the buffer and specify A4 page size
#     doc = SimpleDocTemplate(buffer, pagesize=A4)

#     # Create styles for the Paragraphs
#     styles = getSampleStyleSheet()
#     custom_styles = {
#         'title': ParagraphStyle('Title', parent=styles['Title'], alignment=1, fontSize=18),
#         'text_centered': ParagraphStyle('TextCentered', parent=styles['Normal'], alignment=1),
#         'italic': ParagraphStyle('Italic', parent=styles['Normal'], fontName='Times-Italic'),
#         'bullet': ParagraphStyle('Bullet', parent=styles['Bullet'], spaceBefore=10),
#         'bold': ParagraphStyle('Bold', parent=getSampleStyleSheet()['Normal'], fontName='Times-Bold', fontSize=14, leading=30),
#         'title': ParagraphStyle('Title', parent=styles['Title'], alignment=1, fontSize=18),
#         'subheading': ParagraphStyle('Subheading', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12, spaceBefore=10),
#         'normal': ParagraphStyle('Normal', parent=styles['Normal'], fontSize=12, spaceBefore=6),

#     }

#     # Create a list to store the content
#     content = []

#     # Add the title and contact information
#     content.append(Paragraph(f"<b>{user_obj.first_name} {user_obj.last_name}</b>", custom_styles['title']))
#     content.append(Paragraph(f"{user_obj.full_name} • {cv_obj.email}", custom_styles['text_centered']))
#     # content.append(Spacer(1, 10, color=colors.black, width=600, height=1))  # Line separator

#     # Add the personal statement
#     content.append(Paragraph("<b>PERSONAL STATEMENT</b>", custom_styles['bold']))
#     content.append(Paragraph("Students may use a short statement summarising the purpose of the CV", custom_styles['italic']))
#     content.append(Spacer(1, 10))  # Add some space

#     # Add skills and qualities
#     content.append(Paragraph("<b>SKILLS AND QUALITIES</b>", custom_styles['bold']))
#     # Add skills using bullet points
#     # skills = ["Skill 1 description attached.", "Skill 2 description attached.", "Skill 3 description attached."]
#     skills= skill_obj
#     for skill in skills:
#         content.append(Paragraph(f"• {skill.get_skill_dropdown_display()}", custom_styles['bullet']))

#     # Add qualities using bullet points
#     # qualities = ["Quality 1 description attached.", "Quality 2 description attached.", "Quality 3 description attached."]
#     qualities=quality_obj
#     for quality in qualities:
#         content.append(Paragraph(f"• {quality.get_quality_dropdown_display()}", custom_styles['bullet']))
    
#     content.append(Paragraph("<b>EDUCATION</b>", custom_styles['bold']))

#     # Create a table for Junior Cert
#     junior_cert_table_data = [
#         # ["MONTH YEAR", "JUNIOR CERT", "SCHOOL"],
#         ["Subject", "Grade"],
#         [f"{junioir_cert_obj.subject} ", f"{junioir_cert_obj.result}"],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         # Add more rows as needed
#     ]
#     junior_cert_table = Table(junior_cert_table_data, colWidths=[80, 200, 200])
#     junior_cert_table.setStyle(TableStyle([
#         ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
#         ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
#         ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
#         ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
#         ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
#         ('BACKGROUND', (0, 1), (-1, 1), colors.beige),
#     ]))
#     content.append(junior_cert_table)
    
#     # Create a table for Leaving Cert
#     leaving_cert_table_data = [
#         # ["MONTH YEAR", "LEAVING CERT", "SCHOOL"],
#         ["Subject", "Grade"],
#         [f"{leave_cert_obj.subject}", f"{leave_cert_obj.result}"],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         [" ", " "],
#         # Add more rows as needed
#     ]
#     leaving_cert_table = Table(leaving_cert_table_data, colWidths=[80, 200, 200])
#     leaving_cert_table.setStyle(TableStyle([
#         ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
#         ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
#         ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
#         ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
#         ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
#         ('BACKGROUND', (0, 1), (-1, 1), colors.beige),
#     ]))
#     content.append(leaving_cert_table)

#     content.append(Paragraph("WORK EXPERIENCE", custom_styles['subheading']))

#     # Add job details
#     # job_details = [
#     #     ("DATES FROM – TO (MOST RECENT)", "JOB TITLE, COMPANY, ADDRESS."),
#     #     ("Describe your responsibilities and achievements in terms of impact and results. Use examples, but keep it short.", "DATES FROM – TO"),
#     #     ("JOB TITLE, COMPANY", "Describe your responsibilities and achievements in terms of impact and results. Use examples, but keep it short."),
#     # ]
#     job_details= exp_obj

#     for job_detail in job_details:
#         content.append(Paragraph(job_detail.job_title, custom_styles['normal']))
#         content.append(Paragraph(job_detail.company, custom_styles['normal']))
#         content.append(Spacer(1, 10)) 


#     # Build the content and save it to the buffer
#     doc.build(content)

#     # Move the buffer's file pointer to the beginning
#     buffer.seek(0)

#     return buffer

class GeneratePDF(CreateAPIView):
    def get(self, request):
        """Fetch All Notes By Officer"""
        try:
            print("work")
            student =self.request.user
            user_obj=Student.objects.get(id=student.student.id)
            cv_obj =CV.objects.filter(user=student.student).first()
            education_obj=Education.objects.filter(user=student.student)
            junior_cert_obj=JuniorCertTest.objects.filter(user=student.student)
            leave_cert_obj=LeavingCertTest.objects.filter(user=student.student)
            exp_obj=Experience.objects.filter(user=student.student)
            skill_obj=Skills.objects.filter(user=student.student)
            quality_obj=Qualities.objects.filter(user=student.student)
            interest_obj=Interests.objects.filter(user=student.student)
            additional_info = AdditionalInfo.objects.filter(user=student.student)
            add_info = [i.additional_info for i in additional_info][0]
            refer_obj=Reference.objects.filter(user=student.student)
            temp_name = "general/templates/"
            cv_template = str(user_obj.first_name) +"-"+str(user_obj.last_name) +"-"+"cv" + ".html"
            full_name = user_obj.full_name
            if cv_obj.number == "" or cv_obj.number is None:
                number = "Phone No."
            else:
                number = cv_obj.number
            if full_name:
                words = full_name.split()
                print(len(words))
                if len(words) == 3:
                    print(words)
                    first_name = words[0] + " " + words[1]
                    last_name = words[2]
                elif len(words) == 2:
                    print(words, "working")
                    first_name = words[0]
                    last_name = words[1]
                elif len(words) == 1:
                    first_name = words[0]
                    last_name = ""
                else:
                    first_name = ""
                    last_name = ""
            print(first_name, last_name)
            context = {
                'first_name': first_name,
                'number': number,
                'last_name': last_name,
                'student_detail': user_obj,
                'cv_detail': cv_obj,
                'education_detail': education_obj,
                'Junior_Cert_detail': junior_cert_obj,
                'Leave_Cert_detail': leave_cert_obj,
                'skill_detail': skill_obj,
                'qualities_detail': quality_obj,
                'Experience_detail': exp_obj,
                'Interest_detail':interest_obj,
                'additional_info': add_info,
                'Reference_detail': refer_obj,
                }
            #   open(temp_name + cv_template, "w").write(render_to_string('cv.html',context))
            rendered_template = render_to_string('cv.html', context)
            open(temp_name + cv_template, "w").write(rendered_template)
            HTML(temp_name + cv_template).write_pdf(str("cv")+'.pdf')
            file_location = f'{"cv"}.pdf'
            with open(file_location, 'rb') as f:
                file_data = f.read()
            response = HttpResponse(file_data, content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="'+ "cv" +'".pdf'
            os.remove(temp_name + cv_template)
            os.remove("cv.pdf")
            return response
        except Exception as e:
            print(e)
            return Response({'message': "All steps of CV should be completed"}, status=status.HTTP_400_BAD_REQUEST)


class GenerateAndSendPDF(CreateAPIView):
    def post(self, request):
        """Generate PDF, Send Email with Attachment, and Delete Files"""
        try:
            receiver_email = request.data.get("email")
            student = self.request.user
            user_obj = Student.objects.get(id=student.student.id)
            cv_obj =CV.objects.filter(user=student.student).first()
            education_obj=Education.objects.filter(user=student.student)
            junior_cert_obj=JuniorCertTest.objects.filter(user=student.student)
            leave_cert_obj=LeavingCertTest.objects.filter(user=student.student)
            exp_obj=Experience.objects.filter(user=student.student)
            skill_obj=Skills.objects.filter(user=student.student)
            quality_obj=Qualities.objects.filter(user=student.student)
            interest_obj=Interests.objects.filter(user=student.student)
            additional_info = AdditionalInfo.objects.filter(user=student.student)
            add_info = [i.additional_info for i in additional_info][0]
            refer_obj=Reference.objects.filter(user=student.student)

            temp_name = "general/templates/"
            cv_template = str(user_obj.first_name) + "-" + str(user_obj.last_name) + "-" + "cv" + ".html"
            pdf_filename = str(user_obj.first_name) + ".pdf"
            full_name = user_obj.full_name
            if cv_obj.number == "" or cv_obj.number is None:
                number = "Phone No."
            else:
                number = cv_obj.number
            if full_name:
                words = full_name.split()
                print(len(words))
                if len(words) == 3:
                    print(words)
                    first_name = words[0] + " " + words[1]
                    last_name = words[2]
                elif len(words) == 2:
                    print(words, "working")
                    first_name = words[0]
                    last_name = words[1]
                elif len(words) == 1:
                    first_name = words[0]
                    last_name = ""
                else:
                    first_name = ""
                    last_name = ""
            print(first_name, last_name)
            context = {
                'first_name': first_name,
                'number': number,
                'last_name': last_name,
                'student_detail': user_obj,
                'cv_detail': cv_obj,
                'education_detail': education_obj,
                'Junior_Cert_detail': junior_cert_obj,
                'Leave_Cert_detail': leave_cert_obj,
                'skill_detail': skill_obj,
                'qualities_detail': quality_obj,
                'Experience_detail': exp_obj,
                'Interest_detail':interest_obj,
                'additional_info': add_info,
                'Reference_detail': refer_obj,
            }

            rendered_template = render_to_string('cv.html', context)
            open(temp_name + cv_template, "w").write(rendered_template)
            HTML(temp_name + cv_template).write_pdf(pdf_filename)

            with open(pdf_filename, 'rb') as f:
                file_data = f.read()

            try:
                email = EmailMessage(
                    'Your CV PDF',
                    'Please find your CV PDF attached.',
                    f"{os.environ['EMAIL_HOST_USER']}",  # Replace with your email address
                    [receiver_email],  # List of recipient email addresses
                )
                email.attach(f'{user_obj.first_name}.pdf', file_data, 'application/pdf')

                email.send()
            except Exception as e:
                return Response({"message": f"Error sending email: {e}"})

            # Delete the generated files
            os.remove(temp_name + cv_template)
            os.remove(pdf_filename)

            return Response({'message': "CV PDF sent successfully"}, status=status.HTTP_200_OK)

        except Exception as e:
            return Response({'message': "All steps of CV should be completed"}, status=status.HTTP_400_BAD_REQUEST)


class GenerateDOCX(CreateAPIView):
    def get(self, request):
        """Fetch All Notes By Officer"""
        try:
            # print("work")
            student =self.request.user
            user_obj=Student.objects.get(id=student.student.id)
            cv_obj =CV.objects.filter(user=student.student).first()
            education_obj=Education.objects.filter(user=student.student)
            junior_cert_obj=JuniorCertTest.objects.filter(user=student.student)
            leave_cert_obj=LeavingCertTest.objects.filter(user=student.student)
            exp_obj=Experience.objects.filter(user=student.student)
            skill_obj=Skills.objects.filter(user=student.student)
            quality_obj=Qualities.objects.filter(user=student.student)
            interest_obj=Interests.objects.filter(user=student.student)
            additional_info = AdditionalInfo.objects.filter(user=student.student)
            add_info = [i.additional_info for i in additional_info][0]
            refer_obj=Reference.objects.filter(user=student.student)
            full_name = user_obj.full_name
            if cv_obj.number == "" or cv_obj.number is None:
                number = "Phone No."
            else:
                number = cv_obj.number
            if full_name:
                words = full_name.split()
                if len(words) == 3:
                    first_name = words[0] + " " + words[1]
                    last_name = words[2]
                elif len(words) == 2:
                    first_name = words[0]
                    last_name = words[1]
                elif len(words) == 1:
                    first_name = words[0]
                    last_name = ""
                else:
                    first_name = ""
                    last_name = ""

            doc = Document()
            
            name = doc.add_heading(f'{first_name} {last_name}')
            name.alignment = 1
            name.runs[0].font.size = Pt(32)
            name.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph(f'{cv_obj.email} - {number} - {cv_obj.address} - {cv_obj.eircode}')
            doc.add_paragraph('_' * 100)

            doc.add_heading('PERSONAL STATEMENT').runs[0].font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph(f'{cv_obj.objective}')

            skill_and_quality = doc.add_heading('SKILLS AND QUALITIES', level=1)
            skill_and_quality.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            doc.add_heading('Skills:', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)

            #Skills data
            skills_data = doc.add_paragraph()
            for skill in skill_obj:
                run = skills_data.add_run(f'{skill.skill_dropdown_value}\n')
                run.bold = True
                skills_data.add_run(f'{skill.description}\n')

            #Qualities data                
            doc.add_heading('Qualities:', level=2).runs[0].font.color.rgb = RGBColor(0, 0, 0)
            qualities_data = doc.add_paragraph()
            for quality in quality_obj:
                run = qualities_data.add_run(f'{quality.quality_dropdown_value}\n')
                run.bold = True
                qualities_data.add_run(f'{quality.description}\n')

            #Education data
            education = doc.add_heading('EDUCATION', level=1)
            education.runs[0].font.color.rgb = RGBColor(0, 0, 0)
            for edu in education_obj:
                edu_duration = doc.add_paragraph().add_run(f"""{edu.year} - {'I am still studying here' if edu.present else str(edu.enddate)}""")
                edu_duration.bold = True
                edu_data = doc.add_paragraph().add_run(f"{edu.school}, {edu.examtaken}")

            #Junior Cert
            if len(junior_cert_obj) != 0:
                doc.add_heading('Junior Cert', level=2).runs[0].font.color.rgb = RGBColor(41, 136, 85)

                # Add a table
                table = doc.add_table(rows=len(junior_cert_obj) + 1, cols=3)

                # Add headings
                table.cell(0, 0).text = 'subject'
                table.cell(0, 1).text = 'Level'
                table.cell(0, 2).text = 'Result'

                # Add data in table
                for row_num, cert_data in enumerate(junior_cert_obj, start=1):
                    table.cell(row_num, 0).text = cert_data.subject
                    table.cell(row_num, 1).text = cert_data.level
                    table.cell(row_num, 2).text = cert_data.result

            #Leaving Cert
            if len(leave_cert_obj) != 0:
                doc.add_heading('Leaving Cert', level=2).runs[0].font.color.rgb = RGBColor(41, 136, 85)

                # Add a table
                table = doc.add_table(rows=len(leave_cert_obj) + 1, cols=3)

                # Add headings
                table.cell(0, 0).text = 'subject'
                table.cell(0, 1).text = 'Level'
                table.cell(0, 2).text = 'Result'

                # Add data in table
                for row_num, cert_data in enumerate(leave_cert_obj, start=1):
                    table.cell(row_num, 0).text = cert_data.subject
                    table.cell(row_num, 1).text = cert_data.level
                    table.cell(row_num, 2).text = cert_data.result      

            #Work Experience
            doc.add_heading('WORK EXPERIENCE', level=1).runs[0].font.color.rgb = RGBColor(0, 0, 0)
            for exp in exp_obj:
                exp_duration = doc.add_paragraph().add_run(f"""{exp.startdate} - {'I am currently working here' if exp.is_current_work else str(exp.enddate)}""")
                exp_duration.bold = True
                heading_paragraph = doc.add_paragraph()
                job_title_run = heading_paragraph.add_run(exp.job_title)
                job_title_run.bold = True
                job_title_run.font.color.rgb = RGBColor(41, 136, 85)
                heading_paragraph.add_run(f', {exp.company}, {exp.city}.')
                doc.add_paragraph(exp.description)

            #ACHIEVEMENTS
            doc.add_heading('ACHIEVEMENTS', level=1).runs[0].font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph(add_info)

            #HOBBIES AND INTERESTS
            doc.add_heading('HOBBIES AND INTERESTS', level=1).runs[0].font.color.rgb = RGBColor(0, 0, 0)
            for interest in interest_obj:
                doc.add_paragraph(f'{interest.interests}, {interest.description}')

            #REFEREES
            doc.add_heading('REFEREES', level=1).runs[0].font.color.rgb = RGBColor(0, 0, 0)
            for ref in refer_obj:
                doc.add_paragraph(f'{ref.name}, {ref.position}')
                doc.add_paragraph(f'{ref.contact_number}, {ref.email}')

            doc_io = BytesIO()
            doc.save(doc_io)

            doc_io.seek(0)

            response = HttpResponse(
                doc_io.getvalue(), 
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

            response['Content-Disposition'] = f'attachment; filename={full_name}.docx'

            return response
        
        except Exception as e:
            return Response({'message': str(e) + " All steps of CV should be completed"}, status=status.HTTP_400_BAD_REQUEST)
